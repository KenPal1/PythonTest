import base64, re
import os
import hashlib
import pythoncom
import json
from win32com.client import Dispatch
from io import BytesIO
from PIL import Image
from django.core.files.base import ContentFile
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.forms import AuthenticationForm
from django.contrib.auth.hashers import make_password
from django.contrib.auth.decorators import user_passes_test, login_required
from django.contrib.sessions.backends.db import SessionStore
from django.db.models import Sum, Count, Q, Prefetch
from webapp.models import Appointment, Payment, CustomUser, ServiceType, Patient, Examination
from django.utils import timezone
from django.utils.crypto import get_random_string
from django.http import HttpResponseRedirect, JsonResponse, FileResponse, HttpResponse
from django.urls import reverse
from datetime import datetime
from .forms import UserCreationForm, EditAccountForm, EditProfileForm, AppointmentForm, ExaminationForm, UploadEditedDocumentForm, UploadResultImageForm, EditExaminationForm
from django.views.decorators.csrf import csrf_protect, csrf_exempt
from docx import Document
from docx.shared import Inches
from django.conf import  settings
from django.core.exceptions import ObjectDoesNotExist


def login_view(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(request, username=username, password=password)

        if user is not None:
            login(request, user)
            
            # Determine user type and redirect accordingly
            if user.is_superuser:  # Assuming superuser is the admin
                return redirect('admin_dashboard')
            elif hasattr(user, 'is_employee') and user.is_employee:
                return redirect('employee_dashboard')
            elif hasattr(user, 'is_associated_doctor') and user.is_associated_doctor:
                return redirect('assoc_doc_dashboard')
            else:
                # Redirect to a default page or show an error
                messages.error(request, "Access not allowed.")
                return redirect('login')  # Replace 'login' with your login page name
        else:
            messages.error(request, "Invalid credentials")
            
    return render(request, 'login.html')

@csrf_protect
def admin_login_view(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)

        if user is not None and user.is_superuser:
            login(request, user)
            # Set custom session data (Django manages session storage automatically)
            session_key = f"admin_session_{get_random_string(32)}"
            request.session['session_key'] = session_key
            request.session['user_id'] = user.id
            response = HttpResponseRedirect(reverse('admin_dashboard'))
            response.set_cookie('admin_session', session_key, httponly=True)
            return response
        else:
            messages.error(request, "Invalid credentials or unauthorized access.")
    return render(request, 'admin/admin_login.html')


def employee_login_view(request):
    if request.method == 'POST':
        form = AuthenticationForm(request, data=request.POST)
        if form.is_valid():
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password')
            user = authenticate(request, username=username, password=password)

            if user is not None and hasattr(user, 'is_employee') and user.is_employee:
                login(request, user)
                # Set custom session data (Django manages session storage automatically)

                response = HttpResponseRedirect(reverse('employee_dashboard'))

                return response
            else:
                messages.error(request, "Invalid credentials or unauthorized access.")
    else:
        form = AuthenticationForm()

    return render(request, 'employee/employee_login.html', {'form': form})


@user_passes_test(lambda u: u.is_superuser)
def manage_account_view(request):
    # Filter out the superuser accounts and get only employees and associated doctors
    accounts = CustomUser.objects.exclude(is_superuser=True).filter(Q(is_employee=True) | Q(is_associated_doctor=True) | Q(is_clinic_doctor=True)
    )

    # Context to pass to the template
    context = {
        'accounts': accounts,
    }
    return render(request, 'admin/manage_accounts.html', context)

@user_passes_test(lambda u: u.is_superuser)
def patients_list_view(request):
    return render(request, 'admin/patients_list.html')

@user_passes_test(lambda u: u.is_superuser)
def admin_dashboard_view(request):
    # today = timezone.now()
    # start_of_week = today - timedelta(days=today.weekday())

    # # Total income by payment method
    # daily_income_cash = Payment.objects.filter(date__date=today.date(), method='cash').aggregate(total=Sum('amount'))['total'] or 0
    # daily_income_gcash = Payment.objects.filter(date__date=today.date(), method='gcash').aggregate(total=Sum('amount'))['total'] or 0
    # weekly_income_cash = Payment.objects.filter(date__gte=start_of_week, method='cash').aggregate(total=Sum('amount'))['total'] or 0
    # weekly_income_gcash = Payment.objects.filter(date__gte=start_of_week, method='gcash').aggregate(total=Sum('amount'))['total'] or 0

    # # Count total patients attended
    # # Count total patients for today
    # today_patients = Patient.objects.count

    # # Count total patients for the week
    # weekly_patients = Patient.objects.count

    # # Get today's appointments
    # today_appointments = Appointment.objects.filter(date__date=today.date())
    
    # # Get all appointments
    # all_appointments = Appointment.objects.all()

    # context = {
    #     'daily_income_cash': daily_income_cash,
    #     'daily_income_gcash': daily_income_gcash,
    #     'weekly_income_cash': weekly_income_cash,
    #     'weekly_income_gcash': weekly_income_gcash,
    #     'today_patients': today_patients,  # Total distinct patients today
    #     'weekly_patients': weekly_patients,  # Total distinct patients this week
    #     'today_appointments': today_appointments,  # Today's appointments
    #     'all_appointments': all_appointments,  # All appointments
    # }

    return render(request, 'admin/admin_dashboard.html')

def admin_logout_view(request):
    logout(request)  # This logs out the user
    return redirect('admin_login') 

def employee_logout_view(request):
    logout(request)  # This logs out the user
    return redirect('employee_login') 


def create_account_view(request):
    if request.method == 'POST':
        form = UserCreationForm(request.POST, request.FILES)  # Include `request.FILES` for file uploads
        if form.is_valid():
            # Save the new user
            user = form.save(commit=False)
            user.username = form.cleaned_data['email']
            user.password = make_password(form.cleaned_data['password'])
            user.first_name = form.cleaned_data['first_name']
            user.last_name = form.cleaned_data['last_name']
            user.email = form.cleaned_data['email']

            # Determine account type
            account_type = request.POST.get('account_type')  # Get the selected account type
            if account_type == 'employee':
                user.is_employee = True
            elif account_type == 'assoc_doctor':
                user.is_associated_doctor = True
                user.signature_image = form.cleaned_data.get('signature_image')  # Add signature image
            elif account_type == 'clinic_doctor':
                # Check if a clinic doctor already exists
                if CustomUser.objects.filter(is_clinic_doctor=True).exists():
                    messages.error(request, "A clinic doctor account already exists.")
                    return redirect('create_account')  # Redirect back to the form
                user.is_clinic_doctor = True
                user.signature_image = form.cleaned_data.get('signature_image')  # Add signature image

            # Save the user to the database
            user.save()

            messages.success(request, "Account created successfully!")
            return redirect('admin_dashboard')
        else:
            messages.error(request, "Error creating account. Please check the details.")

    else:
        form = UserCreationForm()
    return render(request, 'admin/admin_dashboard.html', {'form': form})

@user_passes_test(lambda u: u.is_employee)
def employee_dashboard_view(request):
    account = request.user
    # Fetch all appointments and service types for the dashboard
    appointments = Appointment.objects.all()  # or any filtering needed
    service_types = ServiceType.objects.all()

    if request.method == 'POST' and 'add_appointment' in request.POST:
        form = AppointmentForm(request.POST)
        if form.is_valid():
            appointment = form.save(commit=False)
            appointment.save()  # Save the appointment
            form.save_m2m()  # Save many-to-many fields (service types)
            
            # Show success message
            messages.success(request, 'Appointment has been scheduled successfully.')
            return redirect('employee_dashboard')  # Redirect to the dashboard to refresh
        
    else:
        form = AppointmentForm()

    context = {
        'appointments': appointments,
        'service_types': service_types,
        'form': form,
        'account': account,
    }

    return render(request, 'employee/employee_dashboard.html', context)

@login_required
def edit_account_view(request, account_id):
    account = get_object_or_404(CustomUser, id=account_id)
    form = EditAccountForm(request.POST or None, request.FILES or None, instance=account)
    
    if request.method == 'POST':
        if form.is_valid():
            form.save()
            messages.success(request, 'Account updated successfully.')
            return redirect('manage_accounts')
        else:
            messages.error(request, 'Please correct the error below.')
    
    return render(request, 'admin/edit_account.html', {'form': form, 'account': account})

@login_required
def delete_account_view(request, account_id):
    account = get_object_or_404(CustomUser, id=account_id)
    if request.method == 'POST':
        account.delete()
        messages.success(request, 'Account deleted successfully.')
        return redirect('manage_accounts')
    
    return render(request, 'admin/delete_account.html', {'account': account})

@login_required
def edit_profile_view(request, account_id):
    account = get_object_or_404(CustomUser, id=account_id)
    form = EditProfileForm(request.POST or None, request.FILES or None, instance=account)
    
    if request.method == 'POST':
        if form.is_valid():
            form.save()
            messages.success(request, 'Account updated successfully.')
            return redirect('employee_dashboard')
        else:
            messages.error(request, 'Please correct the error below.')
    
    return render(request, 'employee/edit_profile.html', {'form': form, 'account': account})


@user_passes_test(lambda u: u.is_employee)
def employee_patients_list_view(request):
    # Fetching all patients, or you can filter as needed
    patients = Patient.objects.all()
    examinations = (
        Examination.objects.select_related('patient', 'attending_doctor')
        .prefetch_related('service_types', Prefetch('payment_set'))
        .order_by('-date_created')  # Order by latest examination first
    )
    service_types = ServiceType.objects.all()

    account = request.user
    context = {
        'account': account,
        'patients': patients,
        'service_types' : service_types,
        'examinations' : examinations
    }
    return render(request, 'employee/patients_list.html', context)


@user_passes_test(lambda u: u.is_employee)
def document_results_view(request):
       account = request.user
       context = {'account': account}
       return render(request, 'employee/document_results.html', context)


@user_passes_test(lambda u: u.is_employee)
def assoc_doc_readings_view(request):
       account = request.user
       context = {'account': account}
       return render(request, 'employee/assoc_doc_readings.html', context)

@user_passes_test(lambda u: u.is_employee)
def associated_doctors_view(request):
       account = request.user
       associated_doctors = CustomUser.objects.filter(is_associated_doctor=True)
       context = {
           'account': account,
           'associated_doctors': associated_doctors
           }
       
       return render(request, 'employee/associated_doctors.html', context)

def add_appointment(request):
    if request.method == 'POST':
        form = AppointmentForm(request.POST)
        if form.is_valid():
            appointment = form.save(commit=False)
            appointment.save()
            form.save_m2m()
            return redirect('employee_dashboard')  # Redirect back to dashboard after saving
    else:
        form = AppointmentForm()

    # Get the list of service types to display in the modal
    service_types = ServiceType.objects.all()

    context = {
        'form': form,
        'service_types': service_types,
    }
    return render(request, 'employee/add_appointment.html', context)

@user_passes_test(lambda u: u.is_employee)
def employee_examination_view(request):
    # Fetch examinations along with related patient and payment details
    account = request.user
    examinations = (
        Examination.objects.select_related('patient', 'attending_doctor')
        .prefetch_related('service_types', Prefetch('payment_set'))
        .order_by('-date_created')  # Order by latest examination first
    )
    service_types = ServiceType.objects.all()

    context = {
        'account' : account,
        'examinations': examinations,
        'service_types': service_types
    }
    return render(request, 'employee/examination.html', context)


def add_examination(request):
    account = request.user

    if request.method == 'POST':
        # Handle patient search (AJAX request)
        if 'search_patient' in request.POST:
            search_query = request.POST.get('search_patient', '').strip()
            # Filter patients by name or ID
            patients = Patient.objects.filter(
                Q(first_name__istartswith=search_query) |
                Q(last_name__istartswith=search_query) |
                Q(middle_name__istartswith=search_query) |
                Q(id__icontains=search_query)
            )
            # Prepare patient data for JSON response
            patient_list = [
                {
                    'id': patient.id,
                    'first_name': patient.first_name,
                    'last_name': patient.last_name,
                    'middle_name': patient.middle_name or '',
                    'age': patient.age,
                    'sex': patient.sex,
                    'address': patient.address,
                    'contact_number': patient.contact_number,
                    'image_url': patient.image.url if patient.image else None,
                }
                for patient in patients
            ]
            return JsonResponse({'patients': patient_list})

        # Handle form submission for adding an examination
        form = ExaminationForm(request.POST, request.FILES)
        if form.is_valid():
            try:
                # Check if a patient is selected from the search results
                patient_id = request.POST.get('patient_id')
                if patient_id:
                    patient = get_object_or_404(Patient, id=patient_id)
                    
                    # If the patient has an existing image, remove it before adding a new one
                    if patient.image:
                        patient.image.delete()
                    
                    # Handle the new image from the webcam
                    image_data = request.POST.get('image')  # Captured image data from the form
                    if image_data:
                        # Convert base64 image data to a Django ImageFile
                        format, imgstr = image_data.split(';base64,')  # Get base64 string
                        ext = format.split('/')[1]  # Extract file extension (png, jpeg, etc.)
                        image_data = ContentFile(base64.b64decode(imgstr), name=f"patient_{patient.id}_image.{ext}")
                        
                        # Save the new image to the patient instance
                        patient.image = image_data
                        patient.save()
                else:
                    # Create a new patient if not selected
                    first_name = form.cleaned_data['first_name']
                    last_name = form.cleaned_data['last_name']
                    middle_name = form.cleaned_data['middle_name']
                    age = form.cleaned_data['age']
                    sex = form.cleaned_data['sex']
                    address = form.cleaned_data['address']
                    contact_number = form.cleaned_data['contact_number']

                    patient = Patient.objects.create(
                        first_name=first_name,
                        last_name=last_name,
                        middle_name=middle_name,
                        age=age,
                        sex=sex,
                        address=address,
                        contact_number=contact_number,
                    )

                # Create the examination
                examination = Examination.objects.create(
                    patient=patient,
                    attending_doctor=form.cleaned_data['attending_doctor']
                )
                examination.service_types.set(form.cleaned_data['service_types'])

                # Create payment record
                Payment.objects.create(
                    examination=examination,
                    method=form.cleaned_data['method'],
                    amount=form.cleaned_data['amount'],
                    status=form.cleaned_data['status']
                )

                # Generate the document
                generate_examination_document(examination)

                return redirect('employee_examination')  # Redirect to success page
            except Exception as e:
                return render(request, 'employee/add_examination.html', {
                    'form': form,
                    'account': account,
                    'error': f"An error occurred: {e}"
                })
    else:
        form = ExaminationForm()

    return render(request, 'employee/add_examination.html', {
        'form': form,
        'account': account,
    })

def generate_examination_document(examination):
    template_path = 'templates/examination_template.docx'
    output_path = os.path.join(settings.MEDIA_ROOT, 'examination_documents')  # Store the relative path

    if not os.path.exists(output_path):
        os.makedirs(output_path)

    # Load the document template
    doc = Document(template_path)

    # Populate the template with examination details
    patient = examination.patient
    doctor = examination.attending_doctor
    patient_full_name = patient.get_full_name_with_middle_initial()
    doctor_full_name = doctor.get_full_name_with_middle_initial()
    file_number = examination.get_file_number()
    patient_full_name_lastname_starting = patient.patient_full_name_last_name_start()

    # Generate the unique code
    salt = "CHMC2024"
    pepper = "WBMS2025"
    raw_code = f"{file_number}-{patient_full_name}-{doctor_full_name}"
    sha_input = f"{salt}{raw_code}{pepper}".encode('utf-8')
    unique_code = hashlib.sha256(sha_input).hexdigest()[:8].upper()
    unique_code = f"CHMC-{unique_code}"

    # Fill the template with text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = cell.text.replace('{PATIENT_NAME}', patient_full_name)
                cell.text = cell.text.replace('{AGE}', str(patient.age))
                cell.text = cell.text.replace('{SEX}', patient.sex)
                cell.text = cell.text.replace('{SERVICE_TYPE}', ', '.join([str(s) for s in examination.service_types.all()]))
                cell.text = cell.text.replace('{DATE}', examination.date_created.strftime('%B %d, %Y'))
                cell.text = cell.text.replace('{FILE_NO}', file_number)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                # Add patient image if placeholder is found
                    if '{PATIENT_IMAGE}' in cell.text:
                        cell.text = cell.text.replace('{PATIENT_IMAGE}', '')  # Clear the placeholder text
                        if patient.image:
                            run = cell.paragraphs[0].add_run()
                            run.add_picture(patient.image.path, width=Inches(1), height=Inches(1))

    footer = doc.sections[0].footer  # Get the footer from the first section

    # Iterate over paragraphs in the footer
    for paragraph in footer.paragraphs:
        # Replace placeholders in the footer
        if '{DOCTOR_NAME}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{DOCTOR_NAME}', doctor_full_name)

        if '{SIGNATURE}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{SIGNATURE}', '')  # Clear the placeholder
            if doctor.signature_image:
                run = paragraph.add_run()
                run.add_picture(doctor.signature_image.path, width=Inches(1), height=Inches(1))  # Add signature image

        if '{UNIQUE_CODE}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{UNIQUE_CODE}', unique_code)

    # Save the document
    output_filename = f"{patient_full_name_lastname_starting}.docx"
    output_full_path = os.path.join(output_path, output_filename)
    doc.save(output_full_path)

    # Attach the document to the Examination instance with the relative path
    examination.document.name = os.path.relpath(output_full_path, settings.MEDIA_ROOT).replace(os.sep, '/')
    examination.save()


@user_passes_test(lambda u: u.is_employee)
def edit_document(request, examination_id):
    account = request.user
    examination = Examination.objects.get(id=examination_id)
    
    # Get the document's URL. Make sure it is an absolute path (i.e., it starts with http:// or https://)
    document_url = examination.document.url

    # Check if the URL starts with a valid protocol (e.g., 'http://', 'https://')
    if not document_url.startswith('http'):
        document_url = request.build_absolute_uri(examination.document.url)

    return render(request, 'employee/edit_document.html', {
        'account': account,
        'document_url': document_url
    })

def download_document(request, pk):
    examination = get_object_or_404(Examination, pk=pk)
    if examination.document:
        return FileResponse(examination.document.open(), as_attachment=True, filename=examination.document.name)
    return redirect('employee_examination')  # Redirect to the examination list if no document.

def upload_edited_document(request, pk):
    examination = get_object_or_404(Examination, pk=pk)

    if request.method == 'POST':
        form = UploadEditedDocumentForm(request.POST, request.FILES, instance=examination)
        if form.is_valid():
            form.save()
            return redirect('employee_examination')
        else:
            return JsonResponse({'message': 'Failed to upload document.'}, status=400)

    return JsonResponse({'message': 'Invalid request method.'}, status=400)

def docx_to_pdf_exact(docx_path):

    try:
        # Initialize COM library
        pythoncom.CoInitialize()

        # Start MS Word
        word = Dispatch('Word.Application')
        word.Visible = False

        doc = word.Documents.Open(docx_path)
        pdf_path = docx_path.replace('.docx', '.pdf')
        doc.SaveAs(pdf_path, FileFormat=17)  # 17 corresponds to the PDF format
        doc.Close()
        word.Quit()

        # Return the path to the PDF
        return pdf_path
    except Exception as e:
        raise RuntimeError(f"Error during conversion: {e}")
    finally:
        # Ensure COM library is uninitialized
        pythoncom.CoUninitialize()

def view_document(request, pk):

    examination = get_object_or_404(Examination, pk=pk)

    # Ensure an edited document exists
    if not examination.edited_document:
        return redirect('employee_examination')

    # Convert .docx to PDF
    try:
        pdf_path = docx_to_pdf_exact(examination.edited_document.path)

        # Return the PDF file
        return FileResponse(open(pdf_path, 'rb'), content_type='application/pdf')
    except Exception as e:
        # Handle errors gracefully
        return HttpResponse(f"Error during document view: {e}", status=500)
    
def docx_to_pdf_auth_checker(docx_path):
    """
    Converts a .docx file to PDF using MS Word (COM automation).
    """
    try:
        pythoncom.CoInitialize()  # Initialize COM library
        word = Dispatch('Word.Application')
        word.Visible = False  # Keep Word invisible
        doc = word.Documents.Open(docx_path)
        pdf_path = docx_path.replace('.docx', '.pdf')
        doc.SaveAs(pdf_path, FileFormat=17)  # 17 corresponds to PDF format
        doc.Close()
        word.Quit()
        return pdf_path
    except Exception as e:
        raise RuntimeError(f"Error during conversion: {e}")
    finally:
        pythoncom.CoUninitialize()  # Cleanup COM library

def verify_document(request):
    """
    Verify a document by checking the unique code and displaying the PDF if valid.
    """
    if request.method == 'POST':
        unique_code = request.POST.get('unique_code')

        try:
            # Split the code to get necessary parts
            unique_code_parts = unique_code.split('-')
            if len(unique_code_parts) != 2:
                return HttpResponse("Invalid code format.")

            code_prefix, code_value = unique_code_parts
            if code_prefix != "CHMC":
                return HttpResponse("Invalid prefix in the code.")

            # Iterate through all examinations to find the one that matches the unique code
            examination = None
            for ex in Examination.objects.all():
                if ex.get_unique_code() == unique_code:
                    examination = ex
                    break

            if not examination:
                return HttpResponse("Document not found for the provided code.")

            # Check if an edited document exists
            if not examination.has_edited_document():
                return HttpResponse("This document does not have an edited version.")

            # Convert the .docx document to PDF if not already converted
            try:
                pdf_path = docx_to_pdf_auth_checker(examination.edited_document.path)

                # Serve the PDF file for the user to view
                return FileResponse(open(pdf_path, 'rb'), content_type='application/pdf')
            except Exception as e:
                return HttpResponse(f"Error during document conversion: {e}")
        
        except Exception as e:
            return HttpResponse(f"Error: {e}")
    else:
        return render(request, 'authenticity_checker.html')
    
@csrf_exempt
def search_patient(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            query = data.get("query", "").strip()
            if query:
                # Split the query into parts (e.g., "John Doe")
                name_parts = query.split()
                filters = Q()

                if len(name_parts) == 1:
                    # Search for matches in first, middle, or last name
                    filters |= (
                        Q(first_name__icontains=name_parts[0]) |
                        Q(middle_name__icontains=name_parts[0]) |
                        Q(last_name__icontains=name_parts[0])
                    )
                elif len(name_parts) > 1:
                    # Search for first + last name or other combinations
                    filters |= (
                        Q(first_name__icontains=name_parts[0]) & 
                        Q(last_name__icontains=name_parts[1])
                    )
                    # Optionally include middle name if three parts are provided
                    if len(name_parts) > 2:
                        filters |= (
                            Q(first_name__icontains=name_parts[0]) &
                            Q(middle_name__icontains=name_parts[1]) &
                            Q(last_name__icontains=name_parts[2])
                        )

                # Filter patients based on the query
                patients = Patient.objects.filter(filters)[:10]
                results = [
                    {
                        "id": patient.id,
                        "first_name": patient.first_name or "",
                        "middle_name": patient.middle_name or "",
                        "last_name": patient.last_name or "",
                        "age" : patient.age or "",
                        "sex" : patient.sex or "",
                        "contact_number" : patient.contact_number or "",
                        "address" :patient.address or ""
                    }
                    for patient in patients
                ]
                return JsonResponse({"patients": results})

        except json.JSONDecodeError:
            return JsonResponse({"error": "Invalid JSON format"}, status=400)
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)

    return JsonResponse({"patients": []})

def upload_examination_result_image(request, pk):
    if request.method == "POST":
        exam = get_object_or_404(Examination, pk=pk)
        form = UploadResultImageForm(request.POST, instance=exam)
        if form.is_valid():
            image_result = request.POST.get('result_image')
            if image_result:
                        # Convert base64 image data to a Django ImageFile
                format, imgstr = image_result.split(';base64,')  # Get base64 string
                ext = format.split('/')[1]  # Extract file extension (png, jpeg, etc.)
                image_result = ContentFile(base64.b64decode(imgstr), name=f"examination_{exam.id}_image.{ext}")
                        
                        # Save the new image to the patient instance
                exam.result_image = image_result
                exam.save()
                return redirect('employee_examination')
        else:
            return JsonResponse({"status": "error", "message": "Invalid data."})
    return redirect('employee_examination')

def edit_examination(request, pk):
    exam = get_object_or_404(Examination, pk=pk)
    payment = exam.payment_set.first()  # Assuming one payment per exam

    PAYMENT_METHODS = Payment.PAYMENT_METHODS
    if request.method == 'POST':
        form = EditExaminationForm(request.POST)
        if form.is_valid():
            # Update patient details
            if form.cleaned_data.get('patient_first_name'):
                exam.patient.first_name = form.cleaned_data['patient_first_name']
            if form.cleaned_data.get('patient_middle_name'):
                exam.patient.middle_name = form.cleaned_data['patient_middle_name']
            if form.cleaned_data.get('patient_last_name'):
                exam.patient.last_name = form.cleaned_data['patient_last_name']
            exam.patient.save()

            # Update service types
            if form.cleaned_data.get('service_types'):
                exam.service_types.set(form.cleaned_data['service_types'])

            # Update payment details
            if payment:
                payment.method = form.cleaned_data.get('payment_method', payment.method)
                payment.status = form.cleaned_data.get('payment_status', payment.status)
                payment.amount = form.cleaned_data.get('payment_amount', payment.amount)
                payment.save()

            # Use 'next' GET parameter to determine whether to redirect or render
            next_template = request.GET.get('next_template', 'employee/examination.html')  # Default to 'employee/examination.html'
            
            # Check if the next_template is a valid template or return redirect if necessary
            if next_template == 'employee/examination.html':
                return render(request, next_template, {'form': form, 'exam': exam, 'PAYMENT_METHODS' : PAYMENT_METHODS})
            else:
                return redirect(next_template)
        else:
            return JsonResponse({'success': False, 'errors': form.errors})

    # Prepopulate the form for GET requests
    initial_data = {
        'patient_first_name': exam.patient.first_name,
        'patient_middle_name': exam.patient.middle_name,
        'patient_last_name': exam.patient.last_name,
        'service_types': exam.service_types.all(),
        'payment_method': payment.method if payment else '',
        'payment_status': payment.status if payment else '',
        'payment_amount': payment.amount if payment else '',
    }
    form = EditExaminationForm(initial=initial_data)

    # Use the next_template GET parameter to determine which template to render
    next_template = request.GET.get('next_template', 'employee/examination.html')  # Default to 'employee/examination.html'

    return render(request, next_template, {'form': form, 'exam': exam, 'PAYMENT_METHODS' : PAYMENT_METHODS})